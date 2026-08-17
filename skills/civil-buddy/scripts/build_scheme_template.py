"""One-shot generator for references/templates/scheme-cn-a4.docx."""

from __future__ import annotations

import io
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor

from paths import DISCLAIMER, HEADER_STOCK, SKILL_ROOT, TEMPLATE_DEFAULT


def _set_east_asia(run, ascii_name: str, east_asia: str) -> None:
    run.font.name = ascii_name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), ascii_name)
    rFonts.set(qn("w:hAnsi"), ascii_name)
    rFonts.set(qn("w:eastAsia"), east_asia)


def _style_font(style, ascii_name: str, east_asia: str, size_pt: int, bold: bool = False) -> None:
    style.font.name = ascii_name
    style.font.size = Pt(size_pt)
    style.font.bold = bold
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), ascii_name)
    rFonts.set(qn("w:hAnsi"), ascii_name)
    rFonts.set(qn("w:eastAsia"), east_asia)


def _p(doc, text, *, heading=None, size=12, bold=False, center=False, east="宋体", color=None):
    if heading == 1:
        para = doc.add_heading(text, level=1)
    elif heading == 2:
        para = doc.add_heading(text, level=2)
    else:
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        if color is not None:
            run.font.color.rgb = color
        _set_east_asia(run, "Times New Roman", east)
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if heading:
        for run in para.runs:
            run.font.size = Pt(16 if heading == 1 else 14)
            _set_east_asia(run, "Times New Roman", "黑体")
    return para


def build(out: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.left_margin = Mm(25)
    section.right_margin = Mm(25)
    section.top_margin = Mm(25)
    section.bottom_margin = Mm(25)

    _style_font(doc.styles["Normal"], "Times New Roman", "宋体", 12)
    for hid, size in (("Heading 1", 16), ("Heading 2", 14)):
        if hid in doc.styles:
            _style_font(doc.styles[hid], "Times New Roman", "黑体", size, bold=True)

    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hr = hp.add_run(HEADER_STOCK)
    hr.font.size = Pt(9)
    hr.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)
    _set_east_asia(hr, "Times New Roman", "宋体")

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fr = fp.add_run("辖区 {{JURISDICTION}} · {{SHORT_NAME}} · 草稿")
    fr.font.size = Pt(9)
    _set_east_asia(fr, "Times New Roman", "宋体")

    _p(doc, "专项施工方案讨论提纲（AI 草稿）", heading=1, center=True)
    _p(doc, "{{PROJECT_NAME}}", size=18, bold=True, center=True, east="黑体")
    _p(doc, "文件版本：草稿 {{STAMP}}", center=True)
    _p(doc, "辖区：{{JURISDICTION}}", center=True)
    _p(doc, "编制：__________    审核：__________    批准：__________", center=True)

    _p(doc, "2 草稿与责任声明", heading=2)
    _p(doc, DISCLAIMER, size=11)

    _p(doc, "假设与待填", heading=2)
    _p(doc, "{{ASSUMPTIONS}}")

    _p(doc, "3 工程概况", heading=2)
    _p(doc, "{{SEC_OVERVIEW}}")

    _p(doc, "4 编制依据", heading=2)
    _p(doc, "编制依据（已核实）", bold=True, east="黑体")
    _p(doc, "{{CITED_VERIFIED}}")
    _p(doc, "编制依据（未核实）", bold=True, east="黑体")
    _p(doc, "{{CITED_UNVERIFIED}}")

    _p(doc, "5 施工部署与工艺", heading=2)
    _p(doc, "{{SEC_DEPLOY}}")
    _p(doc, "6 质量", heading=2)
    _p(doc, "{{SEC_QUALITY}}")
    _p(doc, "7 安全与应急", heading=2)
    _p(doc, "{{SEC_SAFETY}}")
    _p(doc, "8 环保与文明施工", heading=2)
    _p(doc, "{{SEC_ENV}}")
    _p(doc, "9 资源计划", heading=2)
    _p(doc, "{{SEC_RESOURCES}}")
    _p(doc, "10 验收与资料", heading=2)
    _p(doc, "{{SEC_ACCEPTANCE}}")
    _p(doc, "11 附录", heading=2)
    _p(doc, "{{SEC_APPENDIX}}")

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    _fix_settings(out)


def _fix_settings(docx_path: Path) -> None:
    """python-docx emits <w:zoom w:val="bestFit"/> without required w:percent."""
    buf = io.BytesIO()
    with zipfile.ZipFile(docx_path, "r") as zin:
        with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "word/settings.xml":
                    text = data.decode("utf-8")
                    text = text.replace(
                        '<w:zoom w:val="bestFit"/>',
                        '<w:zoom w:percent="100" w:val="bestFit"/>',
                    )
                    text = text.replace(
                        'w:eastAsia="ja-JP"',
                        'w:eastAsia="zh-CN"',
                    )
                    data = text.encode("utf-8")
                zout.writestr(item, data)
    docx_path.write_bytes(buf.getvalue())


if __name__ == "__main__":
    dest = TEMPLATE_DEFAULT
    build(dest)
    print(dest)
