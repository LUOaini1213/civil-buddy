#!/usr/bin/env python3
"""Smoke: competition submission pack beauty rebuild artifacts.

Drives real generators under output/submission/ and asserts regenerated
PDF/MP4/form paths exist with real project content (not empty template residue).
"""
from __future__ import annotations

import re
import subprocess
import sys
import zipfile
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output" / "submission"


def _run(script: str) -> None:
    p = OUT / script
    assert p.is_file(), p
    r = subprocess.run(
        [sys.executable, str(p)],
        cwd=str(ROOT),
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
        timeout=300,
    )
    print(r.stdout)
    if r.returncode != 0:
        print(r.stderr)
        raise SystemExit(f"FAIL {script} rc={r.returncode}")


def _pdf_text(path: Path) -> str:
    try:
        from pypdf import PdfReader  # type: ignore

        reader = PdfReader(str(path))
        return "\n".join((page.extract_text() or "") for page in reader.pages)
    except Exception:
        pass
    try:
        import PyPDF2  # type: ignore

        reader = PyPDF2.PdfReader(open(path, "rb"))
        return "\n".join((page.extract_text() or "") for page in reader.pages)
    except Exception:
        pass
    # fallback: docx source sibling
    docx = path.with_suffix(".docx")
    if docx.is_file():
        from docx import Document

        d = Document(str(docx))
        parts = [p.text for p in d.paragraphs]
        for t in d.tables:
            for row in t.rows:
                for c in row.cells:
                    parts.append(c.text)
        return "\n".join(parts)
    return ""


def main() -> int:
    assert OUT.is_dir(), OUT
    print("== rebuild docs ==")
    _run("build_submission_docs.py")
    print("== rebuild video ==")
    _run("build_demo_video.py")
    # form refresh after video exists
    _run("build_submission_docs.py")

    pdf1 = OUT / "01-说明文档-装箱拼柜Agent工作台.pdf"
    pdf2 = OUT / "03-人机协同履历表-packing-agent.pdf"
    mp4 = OUT / "02-介绍视频-packing-agent.mp4"
    form = OUT / "00-提交表单填写.txt"
    fails: list[str] = []

    for p in (pdf1, pdf2, mp4, form):
        if not p.is_file() or p.stat().st_size < 500:
            fails.append(f"missing/empty {p.name}")

    t1 = _pdf_text(pdf1)
    if "packing-agent" not in t1 and "装箱拼柜" not in t1:
        fails.append("说明文档 missing title/repo markers")
    if "github.com/LUOaini1213/packing-agent" not in t1 and "LUOaini1213" not in t1:
        # allow docx-only extract
        if "作品链接" not in t1 and "IntentSpec" not in t1:
            fails.append("说明文档 body too sparse / template leftover only")
    bad_tpl = "用简短文字概括作品是做什么的"
    if bad_tpl in t1:
        fails.append("说明文档 still has template instruction phrase")

    t2 = _pdf_text(pdf2)
    if "罗文杰" not in t2 and "L5" not in t2:
        fails.append("履历表 missing name/capability markers")
    if "多智能体" not in t2 and "架构" not in t2:
        fails.append("履历表 missing filled scenario rows")

    # video constraints
    import cv2

    cap = cv2.VideoCapture(str(mp4))
    fps = cap.get(cv2.CAP_PROP_FPS) or 30
    frames = cap.get(cv2.CAP_PROP_FRAME_COUNT) or 0
    cap.release()
    dur = frames / fps if fps else 0
    size_mb = mp4.stat().st_size / 1e6
    print(f"video duration_s={dur:.1f} size_mb={size_mb:.2f}")
    if dur <= 0 or dur > 120:
        fails.append(f"video duration out of range: {dur}")
    if size_mb > 500:
        fails.append(f"video too large: {size_mb}MB")

    # polish markers in video generator source
    vsrc = (OUT / "build_demo_video.py").read_text(encoding="utf-8")
    for m in ("C_PANEL", "C_ACCENT", "rounded", "badge", "grid"):
        if m not in vsrc and m.lower() not in vsrc:
            # rounded_rect or _rounded_rect
            if m == "rounded" and "_rounded_rect" not in vsrc:
                fails.append(f"video generator missing polish marker {m}")
            elif m != "rounded":
                fails.append(f"video generator missing polish marker {m}")

    # form paths resolve
    form_t = form.read_text(encoding="utf-8")
    for name in (pdf1.name, pdf2.name, mp4.name):
        if name not in form_t and str(OUT / name) not in form_t:
            # absolute path may use different separators
            if name.split(".")[0] not in form_t:
                fails.append(f"form missing path marker for {name}")
    # listed absolute paths exist
    for m in re.finditer(r"[A-Za-z]:\\[^\s\r\n]+", form_t):
        p = Path(m.group(0))
        if p.suffix.lower() in {".pdf", ".mp4", ".docx", ".zip", ".txt"} and not p.exists():
            fails.append(f"form path missing on disk: {p}")

    if fails:
        print("FAIL submission_beauty", fails)
        return 1
    print("ALL_PASS submission_beauty")
    print("pdf1_bytes", pdf1.stat().st_size)
    print("pdf2_bytes", pdf2.stat().st_size)
    print("mp4_bytes", mp4.stat().st_size)
    print("duration_s", round(dur, 1))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
