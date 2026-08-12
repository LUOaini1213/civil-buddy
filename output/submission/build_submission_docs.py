# -*- coding: utf-8 -*-
"""Fill competition submission docx from official templates — polished typography."""
from __future__ import annotations

import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Twips

OUT = Path(__file__).resolve().parent
OUT.mkdir(parents=True, exist_ok=True)

TITLE = "装箱拼柜 Agent 工作台 · packing-agent"
REPO = "https://github.com/LUOaini1213/packing-agent"
TEAM = "packing-agent 团队"
MEMBERS = "罗文杰（主控/联调 @LUOaini1213）；崔智（成箱 @cuizhi-chat）；牛东睿（拼柜 @niudongrui）"
SUBMIT_DATE = "2026-08-12"
ABSTRACT = (
    "面向工程出运的多智能体装柜工作台：大Team编排+HITL，小TeamA成箱、小TeamB拼柜；"
    "柜数与坐标由确定性tools计算，LLM不拍N柜不写xyz；含有界辩论与可回放轨迹。"
)

ACCENT = RGBColor(0x2F, 0x5B, 0xC7)
HEADING = RGBColor(0x1A, 0x2B, 0x4A)
BODY = RGBColor(0x2C, 0x33, 0x3A)
MUTED = RGBColor(0x5A, 0x67, 0x7A)


def set_run_font(
    run,
    name: str = "微软雅黑",
    size: float = 11,
    bold: bool = False,
    color: RGBColor | None = None,
) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def clear_para(p) -> None:
    for r in list(p.runs):
        r._element.getparent().remove(r._element)


def set_para(
    p,
    text: str,
    *,
    size: float = 11,
    bold: bool = False,
    name: str = "微软雅黑",
    color: RGBColor | None = None,
    space_after: float = 8,
    space_before: float = 0,
    line_spacing: float = 1.35,
):
    clear_para(p)
    r = p.add_run(text)
    set_run_font(r, name=name, size=size, bold=bold, color=color or BODY)
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing = line_spacing
    return p


def shade_cell(cell, hex_color: str) -> None:
    """Set cell background fill (e.g. 'E8F0FE')."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_margins(cell, top=60, bottom=60, left=80, right=80) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for m, v in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        node = OxmlElement(f"w:{m}")
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    tcPr.append(tcMar)


def cell_set(
    cell,
    text: str,
    size: float = 10,
    *,
    bold: bool = False,
    color: RGBColor | None = None,
    fill: str | None = None,
) -> None:
    for i, para in enumerate(cell.paragraphs):
        if i == 0:
            set_para(
                para,
                text,
                size=size,
                bold=bold,
                color=color or BODY,
                space_after=4,
                space_before=2,
                line_spacing=1.25,
            )
        else:
            clear_para(para)
    if fill:
        shade_cell(cell, fill)
    set_cell_margins(cell)


def style_heading_para(p, level: int = 1) -> None:
    """Polish existing Heading styles after fill."""
    if level == 1:
        size, color, before, after = 16, ACCENT, 16, 10
    else:
        size, color, before, after = 13, HEADING, 12, 6
    # re-style existing runs or whole text
    text = p.text
    if not text.strip():
        return
    set_para(
        p,
        text.strip(),
        size=size,
        bold=True,
        color=color,
        space_before=before,
        space_after=after,
        line_spacing=1.2,
    )


def build_shuoming() -> Path:
    src = Path(r"C:\Users\wenjie.luo\Downloads\说明文档模版.docx")
    dst = OUT / "01-说明文档-装箱拼柜Agent工作台.docx"
    shutil.copy2(src, dst)
    doc = Document(str(dst))

    replacements = {
        "附件1": "附件1 · 说明文档",
        "作品名称": TITLE,
        "作品简介：用简短文字概括作品是做什么的、解决什么问题。（100个字以内）": f"作品简介：{ABSTRACT}",
        "作品链接（如有）：": f"作品链接：{REPO}",
        "团队名称：": f"团队名称：{TEAM}",
        "团队成员：": f"团队成员：{MEMBERS}",
        "提交日期：": f"提交日期：{SUBMIT_DATE}",
        "说明作品面向的主要用户群体。": (
            "主要用户：工厂/工程出运计划员、装箱工艺师、货代操作与审单人员（职场成人，具备物料表与柜型常识）。"
            "行为特征：以 Excel/CSV 物料表、非标尺寸与工期压力为主，需可复核的柜数与装载方案，而非聊天式“随便给个方案”。"
            "数字化能力：会用表格与浏览器；不要求会写代码。验收看 agent_steps 轨迹与 tools 数值，而非黑盒答案。"
        ),
        "建议包含：用户身份、年龄或阶段、行为特征、数字化能力等。": (
            "身份：出运/装箱业务岗与审单岗；阶段：方案编制—确认—拼柜—风险复核。"
            "特征：多票并发、怕“AI编柜数”；要 HITL 闸门、结构/重心可解释、可回放。"
            "能力：浏览器操作 + 表上传即可；高级用户可关简洁演示看组织图与 Agent 输出。"
        ),
        "说明用户在真实场景中遇到的具体问题，以及为什么现有方式不能很好解决。": (
            "痛点1：物料→成箱→订柜→3D装载链路长，人工拆箱/拼柜易漏检重心与进柜。"
            "痛点2：通用 ChatAgent 易“随口说几柜/写坐标”，不可复核、不可出运。"
            "痛点3：纯 3D 装箱软件缺 NL 意图、HITL 闸门与多智能体分工证据。"
            "现有方式：Excel+经验估柜，或孤立装箱算法；不足是难解释、难联调、难证明“tools算数”。"
            "核心问题：在真实出运约束下，交付可观察、可确认、可追责的多智能体装柜方案。"
        ),
        "建议包含：痛点、具体场景、现有解决方式、现有方式存在的不足、数据佐证（如有）、要解决的核心问题。": (
            "场景例：30 模块满载演示停 HITL 确认后拼 1×40HQ，mid50 约 67%；钢件轻量票 mid50 可达 100%。"
            "大票对照（446t 全 Agent 路径）：used=25 柜、mid50≈59%、ship_ok=true、risk=WARN（light 路径仅参考不可单独出运）。"
            "联网校准综合约 9.15/10（harness 0.6.4 / 13 agents，诚实不报本地虚高 9.75 当对外分）。"
        ),
        "描述用户在什么情况下使用智能体，建议列举1-2个场景。": (
            "场景 A · 日常出运：上传/粘贴物料表 → 团队 A 成箱+结构校核 → 人确认柜型与箱方案 → 团队 B 3D 拼柜与 CoG → 风险/裁决。"
            "场景 B · 改方案：在箱表下用自然语言「要一排 / 去掉某材料 / 柜型改 40GP」；可改则重算，不可改返回「无此功能」，不假装成功。"
        ),
        "说明作品能够带来的实际价值，如效率提升、体验改善、成本降低、学习辅助、管理优化等。": (
            "效率：成箱—确认—拼柜一站式，N0* 建议柜数与 3D used 同屏，减少反复对表。"
            "信任：tools 定柜坐标；agent_steps 可回放；HITL 默认停确认闸。"
            "风险：mid50/重心/结构结论可视化；有界辩论抑制无脑加柜。"
            "协作：大 Team⊃A/B 固定专岗，比单 LLM 对话更可分工与验收。"
        ),
        "说明智能体有哪些核心功能，分别解决了什么问题。": (
            "1) 意图与编排：NL/表 → IntentSpec；大 Team 调度 13 专岗。"
            "2) 成箱 TeamA：材料解析、结构半严格校核、装箱方案（两排/一排 snappoint、密装/标准箱）。"
            "3) HITL：确认成箱与柜型后才进拼柜；支持自然语言改方案（可改 / 无此功能契约）。"
            "4) 拼柜 TeamB：N0*、bin3d multi_start、CoG/mid50、风险、三视角+等轴测。"
            "5) 有界辩论 critic↔planner：只改 packing_options，tools 重裁决（非 free swarm）。"
            "6) 可观测：SSE/WebSocket 轨迹、path_honesty、VGM 人签状态、表清洗与 shadow eval。"
        ),
        "建议包含：功能名称、用户输入、处理方式、输出结果、解决的问题等。": (
            "输入：自然语言说明 + 可选 CSV/XLSX 物料表 + 演示预设（满载/钢件）。"
            "处理：steps 主路径固定节点；tools 计算几何与柜数；可选 DeepSeek 仅做意图/解释/影子。"
            "输出：boxes[]、container_plan、mid50、verdict、agent_steps、可视化、PDF/轨迹。"
        ),
        "进一步说明作品具体如何实现": (
            "实现形态：自主开发的 Agent Harness 工作台（非扣子/Dify 模板搭积木）。"
            "前端：Vue2 单页（frontend/index.html），简洁演示默认，闸门与辩论卡可观测。"
            "后端：FastAPI gateway（gateway/app.py）+ packing_assistant harness 0.6.4。"
            "运行时：大 Team 编排小 Team A/B；LangGraph 可选 checkpoint；主路径 agent_mode=steps。"
            "LLM：DeepSeek API 可选；无 Key 时 policy_fallback，steps 仍可用。"
        ),
        "整体流程：用户输入后，智能体如何理解需求、调用知识库/工作流/工具，并生成最终结果；": (
            "流程：NL/表 → IntentSpec → 大Team.orchestrator → 小TeamA（材料/结构/箱方案）→ HITL → "
            "小TeamB（N0*/3D/CoG/评估/风险/可视化）→ 有界 replan（可选）→ finalize。"
            "原则：tools compute numbers; model only routes。禁止 LLM 自由写 xyz/柜数。"
        ),
        "提示词与交互设计：说明是否设计了角色设定、任务指令、输出格式、多轮追问、异常提示等；": (
            "交互：顶栏一键满载演示 + HITL 确认；装箱页自然语言改方案（已应用 / 无此功能）；总览裁决条与风险 pill。"
            "角色：13 固定专岗人设（意图/编排/成箱/HITL/拼柜/风险/收口等），非自由聊天室。"
            "异常：缺尺寸报错条、materials_incomplete 阻断、unsupported 改方案明确提示。"
        ),
        "知识库：如使用知识库，应提交知识库相关材料，包括知识来源、文档清单、知识分类结构、主要内容示例等；": (
            "知识库目录 knowledge_base/ 与 docs/skills/：标准箱库、装载启发式、评分口径、CTU/mid50、结构与出运约束。"
            "skills 示例：material-parse、structure-calc、bin3d-pack、hitl-confirm、replan-critic、risk-cog、vgm-draft。"
            "来源：项目沉淀规则 + 行业习惯软规（两排 1100/1150 snappoint 等），非爬取网页 RAG 堆砌。"
        ),
        "工作流与工具调用：如配置工作流、Skills/技能、插件或外部工具，应提供说明；": (
            "工具注册：tool_registry 分簇 big/A/B；核心 tools 含 packing、structure_calc、booking/N0*、bin3d、cog_repair、"
            "table_mapper、vgm_draft、nl_revision 等。"
            "评测工作流：scripts/workflows/eval-parallel-16.rhai；scripts/test_* 回归门禁。"
        ),
        "如基于平台搭建，说明使用的平台名称（如扣子、Dify），平台能力之外自己做了哪些设计（如有）；如自主开发完成，还需说明采用的前端、后端、数据库、向量知识库、接口设计、部署方式等。": (
            "自主开发。前端 Vue2 静态页；后端 FastAPI+Python；会话/checkpoint 用内存+磁盘/SQLite（LangGraph SqliteSaver）；"
            "知识以文档与规则库为主（非向量库强依赖）。"
            "接口：/api/health、/api/pipeline、/api/team-a、/api/confirm、/api/revise-nl、/api/table/parse、/ws/session/{id} 等。"
            "部署：本机 uvicorn gateway.app:app --host 127.0.0.1 --port 8000；可选 skjolber 3D 服务。"
        ),
        "提供架构图、流程图或思维导图展示作品的整体逻辑、功能模块和使用流程。": (
            "架构（文字图）：NL(+材料) → IntentSpec → [大Team: intent/orchestrator/HITL/critic/finalize]"
            " ├─ 小TeamA: material → structure → box_scheme → present"
            " └─ 小TeamB: planner → loader/bin3d → CoG → eval → risk → visual"
            " Tools 层裁决 N0*/xyz/CoG；Trace: agent_steps + SSE。"
            "详见仓库 docs/ARCHITECTURE.md 与 docs/architecture-as-harness.md。"
        ),
        "说明当前作品的创新点，可从技术、场景、交互、内容组织、用户体验等角度展开。": (
            "1) Tools-first 多智能体：专岗+固定 steps，不是“一个会聊天的装箱 Bot”。"
            "2) HITL 一等公民：默认停确认闸，人确认成箱后再拼柜。"
            "3) 有界辩论 densify-over-raise：抑制无脑加柜，非 free multi-agent swarm。"
            "4) 诚实评测：联网校准约 9.15 与本地 scorecard 分离；path_honesty/VGM 人签/无此功能契约。"
            "5) 领域算法可辩护：N0*、两排 snappoint、mid50/CoG、双口径体积。"
        ),
        "说明当前作品的完成度，存在的不足，以及后续优化方向。": (
            "完成度：可演示端到端产品（网关+UI+13 agents+回归脚本）；核心路径 ship-ready。"
            "不足：TMS/ERP 为 stub；VGM 承运人联调未做；均匀 mid≥0.70 仍弱；大票耗时与 3D 可选服务依赖。"
            "后续：运费/经济性代理深化、更多非标表族、更强可视化与出运签章链路。"
        ),
        "提供作品访问链接，如需要登录，应提供测试账号；如无需登录，可说明直接访问方式。建议附上必要使用步骤，同时提供1-3个测试问题。": (
            f"代码与文档：{REPO}（公开仓库，无需登录）。"
            "本地演示：1) pip install -r requirements.txt "
            "2) uvicorn gateway.app:app --host 127.0.0.1 --port 8000 "
            "3) 浏览器打开 http://127.0.0.1:8000/ （无需账号） "
            "4) 点「满载演示」→ 确认并拼柜 → 看总览/可视化/Agent输出。"
            "测试问题：①柜数是谁算的？②为什么要人确认？③「要一排」会改什么、运费指令为何无此功能？"
        ),
        "展示用户如何使用，建议采用“图片/截图 + 文字说明”的形式。": (
            "案例1 满载 HITL：顶栏满载演示 → 装箱方案出现两排标签与 N0* 建议 → 确认 40HQ 拼柜 → 总览显示 mid50/ship_ok/轨迹。"
            "案例2 自然语言改方案：输入「要一排」→ 状态「已改…」且外宽变单排；输入「帮我算运费」→「无此功能」，箱方案不变。"
            "案例3 表上传：上传 CSV/XLSX → /api/table/parse → 表材料跑 pipeline。"
            "演示脚本详见 docs/competition-demo-script.md；证据页 docs/competition-evidence-one-pager.md。"
        ),
        "说明每位团队成员在作品完成过程中的职责分工和主要贡献。": (
            "罗文杰（@LUOaini1213）：主控编排、harness/gateway、联调发布、评测与比赛交付。"
            "崔智（@cuizhi-chat）：阶段1成箱链路——材料解析、结构计算、装箱方案、知识箱库。"
            "牛东睿（@niudongrui）：阶段2拼柜链路——规划装载、3D/CoG、风险可视化与前端协同。"
            "协作方式：GitHub PR + 固定 API 契约（boxes[]）；AI 辅助编码与文档，数值以 tools 回归为准。"
        ),
    }

    # cover / meta identity styling
    cover_keys = {
        "附件1 · 说明文档": (18, True, ACCENT, 6, 14),
        TITLE: (20, True, HEADING, 4, 10),
        f"作品简介：{ABSTRACT}": (11, False, BODY, 4, 8),
        f"作品链接：{REPO}": (10.5, False, ACCENT, 2, 6),
        f"团队名称：{TEAM}": (11, True, HEADING, 8, 4),
        f"团队成员：{MEMBERS}": (10.5, False, BODY, 2, 4),
        f"提交日期：{SUBMIT_DATE}": (10.5, False, MUTED, 2, 12),
    }

    for p in doc.paragraphs:
        t = p.text.strip()
        style_name = p.style.name if p.style else ""

        if t in replacements:
            new_t = replacements[t]
            if new_t in cover_keys:
                sz, bold, col, sb, sa = cover_keys[new_t]
                set_para(p, new_t, size=sz, bold=bold, color=col, space_before=sb, space_after=sa, line_spacing=1.25)
            elif style_name.startswith("Heading"):
                set_para(p, new_t, size=11, color=BODY, space_after=8, space_before=2, line_spacing=1.35)
            else:
                set_para(p, new_t, size=10.5, color=BODY, space_after=8, space_before=2, line_spacing=1.4)
            continue

        # polish remaining headings that were not replaced
        if style_name == "Heading 1":
            style_heading_para(p, 1)
        elif style_name == "Heading 2":
            style_heading_para(p, 2)
        elif t == "附件1":
            set_para(p, "附件1 · 说明文档", size=18, bold=True, color=ACCENT, space_after=14)
        elif t == "作品名称":
            set_para(p, TITLE, size=20, bold=True, color=HEADING, space_after=10)
        elif t == "团队分工":
            style_heading_para(p, 1)

    # tighten section page margins slightly if possible
    for sec in doc.sections:
        sec.top_margin = Twips(int(720 * 1.1))  # ~0.9"
        sec.bottom_margin = Twips(int(720 * 1.0))
        sec.left_margin = Twips(int(720 * 1.1))
        sec.right_margin = Twips(int(720 * 1.1))

    doc.save(str(dst))
    print("WROTE", dst)
    return dst


def build_resume() -> Path:
    src = Path(r"C:\Users\wenjie.luo\Downloads\人机协同履历表模板.docx")
    dst = OUT / "03-人机协同履历表-packing-agent.docx"
    shutil.copy2(src, dst)
    doc = Document(str(dst))

    # title para
    for p in doc.paragraphs:
        if "人机协同履历表" in (p.text or ""):
            set_para(
                p,
                "人机协同履历表 · packing-agent",
                size=16,
                bold=True,
                color=ACCENT,
                space_after=12,
                space_before=4,
            )
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = doc.tables[0]

    # header row
    r0 = table.rows[0].cells
    cell_set(r0[0], "姓名", size=10.5, bold=True, color=HEADING, fill="E8F0FE")
    cell_set(r0[1], "罗文杰（合并提交负责人）\n协作：崔智、牛东睿", size=10.5, bold=True, fill="F7FAFF")
    cell_set(r0[3], "AI使用时间", size=10.5, bold=True, color=HEADING, fill="E8F0FE")
    cell_set(r0[4], "约2年（含本项目密集开发）", size=10.5, fill="F7FAFF")

    ability = (
        "自我评估：□L1  □L2  □L3  □L4  ☑L5 系统集成\n"
        "能把多智能体、工具链、评测与产品闸门串成可交付系统；并向 L6 partial："
        "制定 tools-first / HITL / 有界辩论规范并约束团队与 AI 协作边界。\n"
        "原则：AI 用于架构讨论、代码生成与文档；柜数/坐标/结构结论一律以仓库 tools 与回归测试为准，禁止模型编造。"
    )
    for c in table.rows[1].cells[1:]:
        cell_set(c, ability, size=9.5, fill="FAFBFD")
    cell_set(table.rows[1].cells[0], "使用能力评估", size=10, bold=True, color=HEADING, fill="E8F0FE")

    # column headers
    r2 = table.rows[2].cells
    cell_set(r2[0], "使用场景", size=10.5, bold=True, color=HEADING, fill="DCE8FC")
    cell_set(r2[1], "AI工具", size=10.5, bold=True, color=HEADING, fill="DCE8FC")
    for c in r2[2:]:
        cell_set(c, "描述（人机分工）", size=10.5, bold=True, color=HEADING, fill="DCE8FC")

    scenarios = [
        (
            "多智能体装柜架构与联调",
            "Cursor / Grok / DeepSeek\n+ 自研 harness",
            "用 AI 辅助拆分大 Team⊃A/B、HITL 与 tools 边界；人工锁定“tools 算数、模型只路由”，并用 test_* 回归验收。",
        ),
        (
            "成箱 / 结构 / 装箱算法迭代",
            "AI 编程助手\n+ Python 回归脚本",
            "对 packing / structure / N0* 等改动先写/跑 scripts 测试；AI 生成草稿，人审几何与两排/一排策略。",
        ),
        (
            "自然语言改方案契约",
            "Grok Build\n+ 单元测试",
            "实现“可改就改 / 不可改返回无此功能”；AI 补解析规则，人以契约测试与 HTTP 验收锁死假成功。",
        ),
        (
            "比赛评测与证据整理",
            "DeepSeek / 联网评测\n+ GitHub",
            "整理联网校准约 9.15、HITL 演示路径与 SCORECARD 边界；AI 起草文档，人校准诚实话术。",
        ),
        (
            "前端演示与简洁模式",
            "AI 辅助前端修改",
            "Vue 演示页 HITL / 辩论卡 / 空态与视觉层级优化；人验证真实点击路径与网关 health。",
        ),
        (
            "知识库与 Skills 沉淀",
            "仓库 docs/skills\n+ AI 摘要",
            "将装载启发式、mid50、VGM 人签等沉淀为可引用文档；AI 帮助归类，人确认领域正确性。",
        ),
    ]
    for i, (scene, tool, desc) in enumerate(scenarios):
        ri = 3 + i
        if ri >= len(table.rows):
            break
        cells = table.rows[ri].cells
        fill = "FFFFFF" if i % 2 == 0 else "F5F8FC"
        cell_set(cells[0], scene, size=9.5, bold=True, color=HEADING, fill=fill)
        cell_set(cells[1], tool, size=9, color=MUTED, fill=fill)
        for c in cells[2:]:
            cell_set(c, desc, size=9, fill=fill)

    doc.save(str(dst))
    print("WROTE", dst)
    return dst


def build_form_and_video_script(doc1: Path, doc2: Path, pdf1: Path | None, pdf2: Path | None, video: Path | None) -> None:
    def abs_p(p: Path | None, fallback: Path) -> str:
        x = p if p and p.exists() else fallback
        return str(x.resolve())

    form = OUT / "00-提交表单填写.txt"
    form.write_text(
        f"""【作品提交 · 直接复制 · 美化版交付】

一、作品标题
{TITLE}

二、作品简介（≤200字）
{ABSTRACT}
{REPO}

三、必填上传文件（绝对路径）
① 说明文档 PDF
   {abs_p(pdf1, OUT / '01-说明文档-装箱拼柜Agent工作台.pdf')}
   源稿：{doc1.resolve()}

② 介绍视频 MP4（≤2分钟）
   {abs_p(video, OUT / '02-介绍视频-packing-agent.mp4')}

③ 人机协同履历表 PDF
   {abs_p(pdf2, OUT / '03-人机协同履历表-packing-agent.pdf')}
   源稿：{doc2.resolve()}

④ 补充资料 ZIP（选填）
   {(OUT / '04-补充资料-packing-agent-supplement.zip').resolve()}

四、演示
本地：uvicorn gateway.app:app --host 127.0.0.1 --port 8000
UI：http://127.0.0.1:8000/
Harness 0.6.4 · 13 agents · 联网校准约 9.15

五、答辩一句话
tools 定柜坐标 · 人确认成箱 · 有界辩论反无脑加柜 · 不是 free swarm
""",
        encoding="utf-8",
    )
    print("WROTE", form)

    script = OUT / "02-介绍视频脚本.md"
    script.write_text(
        f"""# 介绍视频脚本（≤2 分钟）

**成片**：`02-介绍视频-packing-agent.mp4`（美化幻灯讲解 · 自动生成）

作品《{TITLE}》· 开源 {REPO}

口播要点：大Team⊃A/B · tools定柜 · HITL · 有界辩论 · 无此功能契约 · 联网约9.15
""",
        encoding="utf-8",
    )
    print("WROTE", script)


def try_convert_pdf(docx_path: Path) -> Path | None:
    pdf = docx_path.with_suffix(".pdf")
    try:
        import win32com.client  # type: ignore

        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(str(docx_path.resolve()))
        doc.SaveAs(str(pdf.resolve()), FileFormat=17)
        doc.Close(False)
        word.Quit()
        if pdf.exists() and pdf.stat().st_size > 1000:
            print("PDF via Word COM", pdf, "bytes", pdf.stat().st_size)
            return pdf
    except Exception as e:
        print("Word COM failed:", e)

    for cmd in (
        ["soffice", "--headless", "--convert-to", "pdf", "--outdir", str(OUT), str(docx_path)],
        [
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(OUT),
            str(docx_path),
        ],
    ):
        try:
            import subprocess

            subprocess.run(cmd, capture_output=True, text=True, timeout=120)
            if pdf.exists():
                print("PDF via soffice", pdf)
                return pdf
        except Exception as e:
            print("soffice fail", e)
    return None


def build_supplement_zip(doc1: Path, doc2: Path, pdf1: Path | None, pdf2: Path | None) -> Path:
    zpath = OUT / "04-补充资料-packing-agent-supplement.zip"
    root = Path(r"E:\ai比赛")
    files = [
        pdf1,
        pdf2,
        doc1,
        doc2,
        OUT / "00-提交表单填写.txt",
        OUT / "02-介绍视频脚本.md",
        root / "README.md",
        root / "docs" / "ARCHITECTURE.md",
        root / "docs" / "competition-demo-script.md",
        root / "docs" / "competition-evidence-one-pager.md",
        root / "docs" / "research" / "competition-network-review-latest.md",
        root / "docs" / "TEAM.md",
        root / "scripts" / "test_nl_revise_contract.py",
        root / "scripts" / "demo_one_shot.py",
    ]
    with zipfile.ZipFile(zpath, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for f in files:
            if f and Path(f).exists():
                fp = Path(f)
                arc = fp.name if fp.parent == OUT else fp.relative_to(root).as_posix()
                zf.write(fp, arcname=f"packing-agent-supplement/{arc}")
        zf.writestr(
            "packing-agent-supplement/MANIFEST.txt",
            f"repo={REPO}\nbeauty_rebuild=true\n",
        )
    print("WROTE", zpath, "mb", round(zpath.stat().st_size / 1e6, 2))
    return zpath


def main() -> None:
    d1 = build_shuoming()
    d2 = build_resume()
    p1 = try_convert_pdf(d1)
    p2 = try_convert_pdf(d2)
    video = OUT / "02-介绍视频-packing-agent.mp4"
    build_form_and_video_script(d1, d2, p1, p2, video if video.exists() else None)
    build_supplement_zip(d1, d2, p1, p2)
    print("ABSTRACT_LEN", len(ABSTRACT))
    print("DONE OUT", OUT)


if __name__ == "__main__":
    main()
